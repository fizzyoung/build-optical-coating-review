from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SKILL_ROOT = Path(__file__).resolve().parent.parent
OUTPUT = SKILL_ROOT / "assets" / "templates" / "review_outline_template.docx"
PRESET = "compact_reference_guide"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_east_asia_font(style_or_run, name: str = "Microsoft YaHei") -> None:
    style_or_run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def set_style_font(style, latin: str, size: float, color: str | None = None, bold: bool | None = None) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        style.font.bold = bold
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    style._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_bottom_border(paragraph, color: str = "AAB7BE", size: str = "6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_multilevel_heading_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_id = 42
    num_id = 42
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)
    for level, text in enumerate(["%1", "%1.%2", "%1.%2.%3"]):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal")
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "space")
        p_pr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "0")
        ind.set(qn("w:hanging"), "0")
        p_pr.append(ind)
        lvl.extend([start, num_fmt, lvl_text, suff, p_pr])
        abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_heading_number(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(level))
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.insert(0, num_pr)


def set_cell_margins(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS.items():
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_fill(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths: list[int]) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError("Table widths must sum to 9360 DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def format_table_text(table, header: bool = False) -> None:
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(4)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
                    set_east_asia_font(run)
                    if header and row_index == 0:
                        run.bold = True
            if header and row_index == 0:
                set_cell_fill(cell, "E8EEF5")


def add_heading(document: Document, text: str, level: int, num_id: int) -> None:
    paragraph = document.add_paragraph(text, style=f"Heading {level}")
    apply_heading_number(paragraph, num_id, level - 1)


def add_prompt(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    label = paragraph.add_run("写作与证据要求：")
    label.bold = True
    body = paragraph.add_run(text)
    for run in (label, body):
        run.font.name = "Calibri"
        set_east_asia_font(run)
    paragraph.paragraph_format.keep_together = True


def build_document() -> Document:
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    set_style_font(normal, "Calibri", 11, "1F2933")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        1: (16, "2E74B5", 18, 10),
        2: (13, "2E74B5", 14, 7),
        3: (12, "1F4D78", 10, 5),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = document.styles[f"Heading {level}"]
        set_style_font(style, "Calibri", size, color, True)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "Optical Coating Literature Review | Audit Template"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(4)
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("5F6B73")
    add_bottom_border(header)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(4)
    run = footer.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("5F6B73")
    add_page_field(footer)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("红外硫系玻璃基底 DLC 薄膜\n三级大纲与审计写作骨架")
    run.font.name = "Calibri"
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("174A5B")
    set_east_asia_font(run)

    subtitle = document.add_paragraph("材料基础、制备技术与应用进展 | Evidence-traceable review outline")
    subtitle.paragraph_format.space_after = Pt(12)
    for run in subtitle.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor.from_string("5F6B73")
        set_east_asia_font(run)

    meta = document.add_table(rows=6, cols=2)
    meta.style = "Table Grid"
    meta_data = [
        ("项目 ID", "初始化项目时填写；与 project_state.yaml 一致"),
        ("综述类型", "在 Task 02 批准 Narrative Review 或 Systematic Review 后填写"),
        ("大纲版本", "批准版本不得覆盖；新版本通过 SUPERSEDED 关系连接"),
        ("证据门禁", "核心事实 ≥ V3；机制、定量比较与图表数据 ≥ V4"),
        ("Style preset", PRESET),
        ("Named override", "中文字符使用 Microsoft YaHei 作为 East Asian fallback；标题 20 pt #174A5B"),
    ]
    for row, values in zip(meta.rows, meta_data):
        row.cells[0].text, row.cells[1].text = values
        set_cell_fill(row.cells[0], "E8EEF5")
        row.cells[0].paragraphs[0].runs[0].bold = True
    set_table_geometry(meta, [2700, 6660])
    format_table_text(meta)

    document.add_paragraph()
    gate = document.add_table(rows=1, cols=1)
    gate.style = "Table Grid"
    gate.cell(0, 0).text = "状态门禁：本模板只能在 Task 20 的三级大纲获明确批准后用于 Task 21B 写作。每章结束后暂停；未经批准不得继续下一章。"
    set_cell_fill(gate.cell(0, 0), "F4F6F9")
    set_table_geometry(gate, [9360])
    format_table_text(gate)

    num_id = add_multilevel_heading_numbering(document)
    sections = [
        ("引言与综述边界", [
            ("红外窗口的防护需求与应用背景", "把服役环境、透过波段和失效模式连接到涂层需求；用 Source_ID 区分工程背景与直接证据。"),
            ("综述范围、术语和证据方法", "明确硫系玻璃、DLC 类型、工艺、性能、应用、时间与语言边界；说明 Narrative 或 Systematic 分支。")
        ]),
        ("红外硫系玻璃基底的材料基础", [
            ("组成—结构—红外性能关系", "比较玻璃组成、声子吸收、折射率、软化温度与环境稳定性；定量值必须带条件和原文位置。"),
            ("表面加工损伤与镀膜相容性", "综合粗糙度、缺陷、热膨胀、温度上限和清洗敏感性，避免把单一基底结论外推到全部硫系玻璃。")
        ]),
        ("DLC 膜层的结构与光学基础", [
            ("键合结构、缺陷与应力", "区分 sp2/sp3、氢含量、密度和应力的测量与推断；Raman 拟合不得直接等同于精确 sp3 含量。"),
            ("光学常数与红外响应", "在波长、角度、偏振、膜厚和模型条件下比较 n、k、透射、反射和吸收，记录拟合模型与不确定性。")
        ]),
        ("制备技术与工艺窗口", [
            ("PECVD、磁控溅射与离子辅助路线", "比较能量输入、温度、沉积速率、均匀性、应力和基底损伤；把设备差异作为跨研究比较条件。"),
            ("FCVA、PLD 与复合沉积路线", "评估高离化率、颗粒缺陷、过滤、放大和复杂曲面适用性；工程结论需要规模证据。")
        ]),
        ("界面工程与失效控制", [
            ("清洗、活化与过渡层", "追踪表面化学、等离子活化、Si/Ge 等过渡层及梯度设计对附着和光学损耗的影响。"),
            ("残余应力、裂纹与剥落机制", "分别陈述观察结果、作者解释和跨文献判断；因果表述需排除膜厚、温度和缺陷等混杂因素。")
        ]),
        ("综合性能与评价方法", [
            ("光学、机械与环境性能", "按相同测试条件比较透射、硬度、模量、附着、磨损、湿热、盐雾、风沙与热循环，禁止无条件横向排名。"),
            ("表征方法与证据质量", "建立光谱、椭偏、Raman、XPS、SEM、AFM、压痕和划痕的证据矩阵，说明方法局限与互证关系。")
        ]),
        ("应用进展与工程化", [
            ("红外成像、探测与航天窗口", "连接真实器件指标、服役条件和标准要求；实验室样片结果不能直接替代系统级寿命证据。"),
            ("大口径、曲面、均匀性与成本", "综合装夹、温控、过程监控、一致性、良率和产业化成本；明确公开证据不足之处。")
        ]),
        ("共识、争议与研究空白", [
            ("跨文献共识和条件性争议", "按材料、工艺、测试和表征条件解释一致与冲突；所有关系必须连接真实 Source_ID。"),
            ("证据缺口与研究路线图", "从 Claim–Evidence Matrix 和 Literature Map 提炼标准、寿命、界面直接证据及工程放大缺口。")
        ]),
        ("结论", [
            ("材料—工艺—界面—性能综合判断", "只总结正文已建立且达到门禁的主张，保留适用条件、局限和证据强度。"),
            ("未来优先研究方向", "按证据缺口、可验证问题、所需方法和预期工程价值排序，避免空泛展望。")
        ]),
    ]

    for h1, children in sections:
        add_heading(document, h1, 1, num_id)
        for index, (h2, prompt) in enumerate(children, start=1):
            add_heading(document, h2, 2, num_id)
            add_heading(document, "核心命题、证据配置与边界", 3, num_id)
            add_prompt(document, prompt)

    document.add_page_break()
    add_heading(document, "三级标题证据配置表", 1, num_id)
    caption = document.add_paragraph("表前说明：每个三级标题至少连接一个核心命题、一个可核验来源和一个质量门禁。")
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(4)
    caption.paragraph_format.keep_with_next = True
    table = document.add_table(rows=5, cols=4)
    table.style = "Table Grid"
    headers = ["三级标题", "核心命题 / Claim_ID", "Source_ID 与原文位置", "冲突、局限与图表"]
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for row in table.rows[1:]:
        row.cells[0].text = "填写已批准标题"
        row.cells[1].text = "登记命题及最低核验等级"
        row.cells[2].text = "登记来源和页码、章节、图或表"
        row.cells[3].text = "登记冲突证据、边界和追踪 ID"
    set_table_geometry(table, [1800, 2700, 2700, 2160])
    format_table_text(table, header=True)

    add_heading(document, "参考文献", 1, num_id)
    add_prompt(document, "仅纳入正文实际引用且题录已核验的来源；提交版移除内部 ID，但保留正式引文和必要版权声明。")

    core = document.core_properties
    core.title = "红外硫系玻璃基底 DLC 薄膜三级大纲与审计写作骨架"
    core.subject = f"Preset: {PRESET}; evidence-traceable review template"
    core.author = "build-optical-coating-review"
    core.keywords = "DLC, chalcogenide glass, optical coating, literature review, evidence traceability"
    return document


def audit_docx(path: Path) -> None:
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        styles_xml = archive.read("word/styles.xml").decode("utf-8")
        numbering_xml = archive.read("word/numbering.xml").decode("utf-8")
        settings_xml = archive.read("word/settings.xml").decode("utf-8")
    checks = {
        "letter_page": 'w:w="12240"' in document_xml and 'w:h="15840"' in document_xml,
        "one_inch_margins": all(token in document_xml for token in ['w:top="1440"', 'w:right="1440"', 'w:bottom="1440"', 'w:left="1440"']),
        "table_width": 'w:tblW w:type="dxa" w:w="9360"' in document_xml or 'w:tblW w:w="9360" w:type="dxa"' in document_xml,
        "table_indent": 'w:tblInd w:w="120" w:type="dxa"' in document_xml or 'w:tblInd w:type="dxa" w:w="120"' in document_xml,
        "multilevel_numbering": 'w:multiLevelType w:val="multilevel"' in numbering_xml,
        "body_line_spacing": 'w:line="300"' in styles_xml,
        "update_fields": "w:updateFields" in settings_xml,
    }
    failed = [name for name, passed in checks.items() if not passed and name != "update_fields"]
    if failed:
        raise RuntimeError(f"DOCX preset audit failed: {failed}")


def enable_field_updates(document: Document) -> None:
    settings = document.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    enable_field_updates(document)
    document.save(OUTPUT)
    audit_docx(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
